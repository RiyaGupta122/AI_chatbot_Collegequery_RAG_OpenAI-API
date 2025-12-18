import os
import logging
from typing import List, Optional, Union
from pathlib import Path
import magic
from PyPDF2 import PdfReader
from docx import Document
from langchain.docstore.document import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from ..core.config import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process various document types and extract text."""
    
    def __init__(self):
        self.allowed_file_types = settings.ALLOWED_FILE_TYPES
        self.max_file_size = settings.MAX_FILE_SIZE_MB * 1024 * 1024  # Convert MB to bytes
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def is_file_allowed(self, file_path: str) -> bool:
        """Check if the file type is allowed."""
        mime = magic.Magic(mime=True)
        file_type = mime.from_file(file_path)
        
        # Map MIME types to file extensions
        mime_to_extension = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'text/plain': 'txt'
        }
        
        file_extension = mime_to_extension.get(file_type)
        return file_extension in self.allowed_file_types
    
    def is_file_size_valid(self, file_path: str) -> bool:
        """Check if the file size is within limits."""
        return os.path.getsize(file_path) <= self.max_file_size
    
    def process_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            raise ValueError(f"Failed to process PDF: {str(e)}")
    
    def process_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")
    
    def process_txt(self, file_path: str) -> str:
        """Read text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {str(e)}")
            raise ValueError(f"Failed to read text file: {str(e)}")
    
    def process_document(self, file_path: str) -> str:
        """Process document based on file type."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not self.is_file_allowed(file_path):
            raise ValueError(f"File type not allowed: {file_path}")
        
        if not self.is_file_size_valid(file_path):
            raise ValueError(f"File size exceeds maximum allowed size of {settings.MAX_FILE_SIZE_MB}MB")
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.process_pdf(file_path)
        elif file_extension == '.docx':
            return self.process_docx(file_path)
        elif file_extension == '.txt':
            return self.process_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def split_text_into_chunks(self, text: str, metadata: dict = None) -> List[LangchainDocument]:
        """Split text into chunks using the text splitter."""
        if not metadata:
            metadata = {}
            
        # Create a single document first
        doc = LangchainDocument(page_content=text, metadata=metadata)
        
        # Split the document
        return self.text_splitter.split_documents([doc])
    
    def process_files(self, file_paths: List[str], metadata: dict = None) -> List[LangchainDocument]:
        """Process multiple files and return Langchain documents."""
        all_docs = []
        
        for file_path in file_paths:
            try:
                # Process the document
                text = self.process_document(file_path)
                
                # Add file-specific metadata
                file_metadata = {
                    'source': os.path.basename(file_path),
                    'file_path': file_path,
                    'file_type': Path(file_path).suffix.lower(),
                    **(metadata or {})
                }
                
                # Split text into chunks
                chunks = self.split_text_into_chunks(text, file_metadata)
                all_docs.extend(chunks)
                
                logger.info(f"Processed {file_path}: {len(chunks)} chunks created")
                
            except Exception as e:
                logger.error(f"Error processing {file_path}: {str(e)}")
                raise
        
        return all_docs
